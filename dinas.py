# ==============================================================================
# Aplikasi Streamlit untuk Mengisi Laporan Inspeksi dan Rincian Biaya Dinas
# ==============================================================================
import io
import math
import json
from datetime import date, datetime, timedelta
from typing import Dict, Any, List
from urllib.parse import quote_plus

import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
import pandas as pd
from openpyxl import load_workbook
import requests
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from email.mime.text import MIMEText # Tambahkan import ini

# --- MongoDB Imports ---
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from pymongo.errors import ConnectionFailure, OperationFailure

# ==============================================================================
# KONFIGURASI APLIKASI
# ==============================================================================

# MongoDB Configuration
username = quote_plus("laporanglss")
password = quote_plus("Kmzway87aa")
MONGODB_URI = f"mongodb+srv://{username}:{password}@cluster0.fbp5d0n.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"
DATABASE_NAME = "laporan_dinas"
COLLECTION_NAME = "rbd_trips"

# Template URLs
TEMPLATE_INSPEKSI_URL = "https://raw.githubusercontent.com/FajarGLS/Laporan-Dinas/main/INSPEKSI.docx"
TEMPLATE_RBD_URL = "https://raw.githubusercontent.com/FajarGLS/Laporan-Dinas/main/RBD.xlsx"

# Email Configuration
EMAIL_SENDER = "fajar@dpagls.my.id"
# PERBAIKAN: Ganti string ini dengan kata sandi email Anda yang sebenarnya.
# JANGAN gunakan quote_plus() karena itu bukan untuk kata sandi.
EMAIL_PASSWORD = "Rahasia100%"
SMTP_SERVER = "mail.dpagls.my.id"
SMTP_PORT = 465

# Daftar kapal
VESSEL_LISTS = {
    "Bulk Carrier": [
        "MV NAZIHA", "MV AMMAR", "MV NAMEERA", "MV ARFIANIE AYU", "MV SAMI",
        "MV KAREEM", "MV NADHIF", "MV KAYSAN", "MV ABDUL HAMID", "MV NASHALINA",
        "MV GUNALEILA", "MV NUR AWLIYA", "MV NATASCHA", "MV SARAH S", "MV MARIA NASHWAH",
        "MV ZALEHA FITRAT", "MV HAMMADA", "MV KAMADIYA", "MV MUBASYIR", "MV MUHASYIR",
        "MV MUNQIDZ", "MV MUHARRIK", "MV MUMTAZ", "MV UNITAMA LILY", "MV RIMBA EMPAT",
        "MV MUADZ", "MV MUNIF", "MV RAFA", "MV NOUR MUSHTOFA", "MV FEIZA",
        "MV MURSYID", "MV AFKAR", "MV. AMOLONGO EMRAN", "MV. NIMAOME EMRAN",
        "MV. SYABIL EMRAN"
    ],
    "Tanker": [
        "MT BIO EXPRESS", "MT KENCANA EXPRESS", "MT SIL EXPRESS", "MT SELUMA EXPRESS"
    ],
    "Tug&Barge": [
        "TB. AZALEA", "TB. GRENADA", "TB. MAGNOLIA", "TB. ZINNIA", "TB. JAZZY",
        "TB. AMARILLYS", "TB. FENNEL", "TB. LAUREL", "TB. JASMINE", "TB. CAMELIA",
        "TB. MULBERRY", "TB. FUSCHIA", "TB. IRIS", "TB. JUNIPER", "TB. SPEEDWELL",
        "TB. MIMOSA", "TB. IVY", "TB. SORREL", "TB. GERBERA", "TB. CLEMATIS",
        "TB. EUSTOMA", "TB. FEIHA", "TB. EHSAL", "TB. GMS CEMERLANG 1", "TB. GMS CEMERLANG 2",
        "TB. ALYSUM", "TB. CATALEYA", "TB. GMS CEMERLANG 3"
    ]
}

# ==============================================================================
# INISIALISASI SESSION STATE
# ==============================================================================

def init_session_state():
    """Inisialisasi session state untuk semua variabel yang diperlukan"""
    if "report_type" not in st.session_state:
        st.session_state.report_type = "Laporan Inspeksi"
    if "form_data" not in st.session_state:
        st.session_state.form_data = {}
    if "activities" not in st.session_state:
        st.session_state.activities = [{"date": "", "time": "", "description": ""}]

# ==============================================================================
# MONGODB FUNCTIONS
# ==============================================================================

@st.cache_resource
def init_mongodb():
    """Initialize MongoDB connection with caching"""
    try:
        client = MongoClient(MONGODB_URI, server_api=ServerApi('1'), serverSelectionTimeoutMS=5000)
        client.admin.command('ping')
        return client[DATABASE_NAME][COLLECTION_NAME]
    except ConnectionFailure as e:
        st.error(f"❌ Gagal terhubung ke MongoDB. Pastikan IP Whitelisting sudah diatur. Detail: {e}")
        return None
    except OperationFailure as e:
        st.error(f"❌ Gagal terhubung ke MongoDB. Pastikan username dan password benar. Detail: {e}")
        return None
    except Exception as e:
        st.error(f"❌ Gagal terhubung ke MongoDB: {e}")
        return None

def save_to_mongodb(data: Dict[str, Any]):
    """Menyimpan data ke MongoDB"""
    db_collection = init_mongodb()
    if db_collection is None:
        st.error("❌ Koneksi MongoDB tidak tersedia")
        return False
    
    if not data.get("trip_id"):
        st.error("❌ Trip ID tidak boleh kosong untuk menyimpan data.")
        return False
    
    try:
        result = db_collection.replace_one(
            {"trip_id": data["trip_id"]}, 
            data, 
            upsert=True
        )
        
        if result.upserted_id:
            st.success(f"✅ Data perjalanan dinas '{data['trip_id']}' berhasil disimpan!")
        else:
            st.success(f"✅ Data perjalanan dinas '{data['trip_id']}' berhasil diperbarui!")
        return True
    except Exception as e:
        st.error(f"❌ Gagal menyimpan ke MongoDB: {e}")
        return False

def load_from_mongodb(trip_id: str) -> Dict[str, Any]:
    """Memuat data dari MongoDB"""
    db_collection = init_mongodb()
    if db_collection is None:
        st.error("❌ Koneksi MongoDB tidak tersedia")
        return {}
    
    try:
        document = db_collection.find_one({"trip_id": trip_id})
        if document:
            document.pop('_id', None)
            st.success(f"✅ Data perjalanan dinas '{trip_id}' berhasil dimuat!")
            return document
        else:
            st.warning(f"⚠️ Data perjalanan dinas dengan ID '{trip_id}' tidak ditemukan.")
            return {}
    except Exception as e:
        st.error(f"❌ Gagal memuat dari MongoDB: {e}")
        return {}

def get_all_trip_ids() -> list:
    """Mendapatkan semua trip ID yang tersimpan"""
    db_collection = init_mongodb()
    if db_collection is None:
        return []
    
    try:
        trip_ids = [doc["trip_id"] for doc in db_collection.find({}, {"trip_id": 1, "_id": 0})]
        return sorted(trip_ids)
    except Exception as e:
        st.error(f"❌ Gagal mengambil daftar trip ID: {e}")
        return []

# ==============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ==============================================================================

def _replace_in_paragraph(paragraph, placeholder, value):
    """Mengganti placeholder dalam paragraph"""
    if placeholder not in paragraph.text:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text.replace(placeholder, value)
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].clear()
        paragraph.runs[0].text = ""
    if len(paragraph.runs) == 0:
        paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text

def replace_placeholder_everywhere(doc: Document, placeholder: str, value: str):
    """Mengganti placeholder di seluruh dokumen"""
    for p in doc.paragraphs:
        _replace_in_paragraph(p, placeholder, value)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, placeholder, value)

def find_table_with_placeholder(doc: Document, placeholder: str):
    """Mencari tabel yang mengandung placeholder"""
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        return tbl
    return None

def add_activities_to_table(doc: Document, activities: List[Dict]):
    """Menambahkan kegiatan ke tabel di dokumen"""
    # Cari tabel yang mengandung "Hari/Tgl/Jam"
    target_table = None
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if "Hari/Tgl/Jam" in cell.text:
                    target_table = tbl
                    break
            if target_table:
                break
        if target_table:
            break
    
    if not target_table:
        return
    
    # Mulai dari baris kedua (setelah header)
    if len(target_table.rows) < 2:
        return
    
    # Hapus baris kosong yang ada (kecuali header)
    while len(target_table.rows) > 1:
        target_table._tbl.remove(target_table.rows[-1]._element)
    
    # Tambahkan baris untuk setiap aktivitas
    for activity in activities:
        if activity["date"] or activity["time"] or activity["description"]:
            row = target_table.add_row()
            
            # Format tanggal dan waktu
            datetime_str = ""
            if activity["date"]:
                datetime_str += activity["date"]
            if activity["time"]:
                if datetime_str:
                    datetime_str += f" / {activity['time']}"
                else:
                    datetime_str = activity["time"]
            
            row.cells[0].text = datetime_str
            row.cells[1].text = activity["description"]

# ==============================================================================
# EMAIL FUNCTIONS
# ==============================================================================

def send_email_with_attachment(from_email, password, to_email, smtp_server, smtp_port, subject, body, attachments):
    """Mengirim email dengan banyak lampiran"""
    try:
        msg = MIMEMultipart()
        msg["From"] = from_email
        msg["To"] = to_email
        msg["Subject"] = subject
        
        # Perbaikan: Menggunakan MIMEText untuk konten email
        msg.attach(MIMEText(body, 'plain'))

        for attachment_bytes, attachment_filename in attachments:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment_bytes)
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {attachment_filename}",
            )
            msg.attach(part)

        server = smtplib.SMTP_SSL(smtp_server, int(smtp_port))
        server.login(from_email, password)
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()
        return True
    except Exception as e:
        st.error(f"Gagal mengirim email: {e}")
        return False

# ==============================================================================
# CALLBACK FUNCTIONS
# ==============================================================================

def add_activity():
    """Callback untuk menambah aktivitas baru"""
    st.session_state.activities.append({"date": "", "time": "", "description": ""})

def remove_activity(index):
    """Callback untuk menghapus aktivitas"""
    if len(st.session_state.activities) > 1:
        st.session_state.activities.pop(index)

# ==============================================================================
# UI COMPONENTS
# ==============================================================================

def render_sidebar():
    """Render sidebar dengan navigasi dan status"""
    with st.sidebar:
        st.image("https://via.placeholder.com/200x100/0066CC/FFFFFF?text=DPA+GLS", width=200)
        st.markdown("---")
        
        st.header("📋 Pilih Jenis Laporan")
        report_type = st.radio(
            "Pilih laporan yang akan dibuat:",
            ("Laporan Inspeksi", "Rincian Biaya Perjalanan Dinas"),
            key="report_type_radio"
        )
        st.session_state.report_type = report_type
        
        st.markdown("---")
        st.subheader("🗄️ Status Database")
        db_collection = init_mongodb()
        if db_collection is not None:
            st.success("🟢 DB Connected")
        else:
            st.error("🔴 DB Disconnected")
        
        # Info aplikasi
        st.markdown("---")
        st.markdown("""
        **📱 Aplikasi Laporan DPA GLS**
        
        Fitur:
        - 📝 Laporan Inspeksi Kapal
        - 💰 Rincian Biaya Dinas (RBD)
        - 📧 Auto Email
        - 💾 Penyimpanan Cloud
        """)

def render_inspection_form():
    """Render form untuk laporan inspeksi"""
    st.markdown("## 🚢 Laporan Inspeksi Kapal")
    
    # Progress indicator
    progress = st.progress(0)
    progress.progress(10)
    
    # Vessel Details dalam card
    with st.container():
        st.markdown("### 📋 Detail Kapal")
        col1, col2 = st.columns(2)
        
        with col1:
            ship_type = st.selectbox(
                "🏷️ Tipe Kapal", 
                options=list(VESSEL_LISTS.keys()),
                help="Pilih tipe kapal yang akan diinspeksi"
            )
            vessel_options = VESSEL_LISTS.get(ship_type, [])
            vessel_name = st.selectbox(
                "🚢 Nama Kapal", 
                options=vessel_options,
                help="Pilih nama kapal dari daftar"
            )
            imo = st.text_input(
                "🆔 IMO Number", 
                placeholder="Contoh: IMO 1234567",
                help="International Maritime Organization number"
            )
        
        with col2:
            callsign = st.text_input(
                "📻 Call Sign", 
                placeholder="Contoh: ABC123",
                help="Call sign kapal untuk komunikasi radio"
            )
            place = st.text_input(
                "📍 Tempat Inspeksi", 
                placeholder="Contoh: Jakarta",
                help="Lokasi dimana inspeksi dilakukan"
            )
            survey_date = st.date_input(
                "📅 Tanggal Inspeksi", 
                value=date.today(),
                help="Tanggal pelaksanaan inspeksi"
            )
    
    progress.progress(30)
    
    # Personnel Details
    with st.container():
        st.markdown("### 👥 Detail Personnel")
        col1, col2 = st.columns(2)
        
        with col1:
            master = st.text_input(
                "👨‍✈️ Nama Master/Kapten", 
                placeholder="Nama lengkap kapten kapal",
                help="Nama kapten kapal yang bertanggung jawab"
            )
        
        with col2:
            surveyor = st.text_input(
                "🔍 Nama Surveyor", 
                placeholder="Nama surveyor yang melakukan inspeksi",
                value="Mohammad Fajar S",
                help="Nama surveyor DPA yang melakukan inspeksi"
            )
    
    progress.progress(50)
    
    # Activities Section
    with st.container():
        st.markdown("### 📝 Kegiatan Inspeksi")
        st.info("💡 Tambahkan kegiatan-kegiatan yang dilakukan selama inspeksi")
        
        for i, activity in enumerate(st.session_state.activities):
            with st.expander(f"Kegiatan {i+1}", expanded=True):
                col1, col2, col3 = st.columns([2, 2, 1])
                
                with col1:
                    activity["date"] = st.text_input(
                        "📅 Hari/Tanggal", 
                        value=activity["date"],
                        placeholder="Contoh: Senin, 15 Agustus 2024",
                        key=f"date_{i}"
                    )
                
                with col2:
                    activity["time"] = st.text_input(
                        "🕐 Jam", 
                        value=activity["time"],
                        placeholder="Contoh: 08:00 - 10:00",
                        key=f"time_{i}"
                    )
                
                with col3:
                    if st.button("🗑️", help="Hapus kegiatan", key=f"remove_{i}"):
                        remove_activity(i)
                        st.rerun()
                
                activity["description"] = st.text_area(
                    "📋 Deskripsi Pekerjaan", 
                    value=activity["description"],
                    placeholder="Jelaskan detail kegiatan yang dilakukan...",
                    height=100,
                    key=f"desc_{i}"
                )
        
        if st.button("➕ Tambah Kegiatan", type="secondary"):
            add_activity()
            st.rerun()
    
    progress.progress(80)
    
    # Email and Generate Section
    with st.container():
        st.markdown("### 📧 Pengiriman Laporan")
        
        col1, col2 = st.columns(2)
        with col1:
            email_to_send = st.text_input(
                "📬 Email Penerima", 
                placeholder="contoh@email.com",
                help="Email yang akan menerima laporan inspeksi"
            )
        
        with col2:
            st.markdown("#### 📊 Ringkasan")
            st.write(f"🚢 **Kapal:** {vessel_name}")
            st.write(f"📅 **Tanggal:** {survey_date}")
            st.write(f"📝 **Kegiatan:** {len(st.session_state.activities)} item")
    
    progress.progress(100)
    
    # Generate Button
    if st.button("🚀 Generate & Kirim Laporan", type="primary", use_container_width=True):
        generate_inspection_report(
            vessel_name, imo, ship_type, callsign, place, survey_date, 
            master, surveyor, st.session_state.activities, email_to_send
        )

def render_rbd_form():
    """Render form untuk RBD"""
    st.markdown("## 💰 Rincian Biaya Perjalanan Dinas")
    
    # Trip Management
    with st.container():
        st.markdown("### 🆔 Manajemen Perjalanan")
        
        col1, col2 = st.columns([3, 2])
        with col1:
            trip_id = st.text_input(
                "🏷️ ID Perjalanan Dinas", 
                help="Contoh: 'FAJAR01'", 
                value=st.session_state.form_data.get("trip_id", ""),
                placeholder="ID unik untuk perjalanan dinas"
            )
        
        with col2:
            all_trip_ids = get_all_trip_ids()
            if all_trip_ids:
                selected_existing_trip = st.selectbox(
                    "📂 Atau pilih yang sudah ada:", 
                    [""] + all_trip_ids,
                    help="Pilih dari perjalanan dinas yang sudah tersimpan"
                )
                if selected_existing_trip and selected_existing_trip != trip_id:
                    loaded_data = load_from_mongodb(selected_existing_trip)
                    if loaded_data:
                        st.session_state.form_data = loaded_data
                        st.rerun()
    
    # Action buttons
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("💾 Simpan Data", type="secondary"):
            save_rbd_data(trip_id)
    with col2:
        if st.button("📂 Muat Data", type="secondary"):
            if trip_id:
                loaded_data = load_from_mongodb(trip_id)
                if loaded_data:
                    st.session_state.form_data = loaded_data
                    st.rerun()
            else:
                st.warning("⚠️ Silakan masukkan Trip ID untuk memuat data.")
    
    # Trip Details
    with st.container():
        st.markdown("### 📅 Detail Perjalanan")
        
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input(
                "📅 Tanggal Mulai", 
                value=pd.to_datetime(st.session_state.form_data.get("start_date")).date() 
                if st.session_state.form_data.get("start_date") else date.today()
            )
        with col2:
            end_date = st.date_input(
                "📅 Tanggal Selesai", 
                value=pd.to_datetime(st.session_state.form_data.get("end_date")).date() 
                if st.session_state.form_data.get("end_date") else date.today()
            )
        
        trip_purpose = st.text_input(
            "🎯 Tujuan Dinas", 
            help="Akan masuk ke sel C13", 
            value=st.session_state.form_data.get("trip_purpose", ""),
            placeholder="Contoh: Inspeksi Kapal MV NAZIHA"
        )
        
        vessel_code = st.text_input(
            "🚢 Kapal Tujuan Dinas / Vessel Code", 
            help="Akan masuk ke sel F13", 
            value=st.session_state.form_data.get("vessel_code", ""),
            placeholder="Contoh: MV NAZIHA"
        )
    
    # Cost Details
    with st.container():
        st.markdown("### 💵 Rincian Biaya")
        st.info("💡 Masukkan nilai dalam Rupiah (angka saja, tanpa titik atau koma)")
        
        # Create tabs for different cost categories
        tab1, tab2, tab3 = st.tabs(["🏨 Akomodasi", "✈️ Transportasi", "🚕 Lokal Transport"])
        
        with tab1:
            col1, col2 = st.columns(2)
            with col1:
                hotel_cost = st.number_input(
                    "🏨 Akomodasi Hotel", 
                    value=float(st.session_state.form_data.get("hotel_cost", 0)),
                    min_value=0.0,
                    help="Biaya menginap di hotel"
                )
                deposit = st.number_input(
                    "💳 Deposit Hotel", 
                    value=float(st.session_state.form_data.get("deposit", 0)),
                    min_value=0.0,
                    help="Deposit yang dibayarkan ke hotel"
                )
        
        with tab2:
            col1, col2 = st.columns(2)
            with col1:
                plane_cost = st.number_input(
                    "✈️ Pesawat", 
                    value=float(st.session_state.form_data.get("plane_cost", 0)),
                    min_value=0.0
                )
                ship_cost = st.number_input(
                    "🚢 Kapal Laut", 
                    value=float(st.session_state.form_data.get("ship_cost", 0)),
                    min_value=0.0
                )
                train_cost = st.number_input(
                    "🚂 Kereta Api", 
                    value=float(st.session_state.form_data.get("train_cost", 0)),
                    min_value=0.0
                )
            with col2:
                bus_cost = st.number_input(
                    "🚌 Bus", 
                    value=float(st.session_state.form_data.get("bus_cost", 0)),
                    min_value=0.0
                )
                airport_tax = st.number_input(
                    "🏛️ Airport Tax", 
                    value=float(st.session_state.form_data.get("airport_tax", 0)),
                    min_value=0.0
                )
                miscellaneous = st.number_input(
                    "📋 Miscellaneous Document Cargo", 
                    value=float(st.session_state.form_data.get("miscellaneous", 0)),
                    min_value=0.0
                )
        
        with tab3:
            col1, col2 = st.columns(2)
            with col1:
                fuel_cost = st.number_input(
                    "⛽ Kendaraan Dinas (BBM)", 
                    value=float(st.session_state.form_data.get("fuel_cost", 0)),
                    min_value=0.0
                )
                toll_cost = st.number_input(
                    "🛣️ Nota Toll", 
                    value=float(st.session_state.form_data.get("toll_cost", 0)),
                    min_value=0.0
                )
                taxi_cost = st.number_input(
                    "🚕 Taksi / Bus", 
                    value=float(st.session_state.form_data.get("taxi_cost", 0)),
                    min_value=0.0
                )
            with col2:
                local_transport = st.number_input(
                    "🚌 Transportasi di tempat dinas", 
                    value=float(st.session_state.form_data.get("local_transport", 0)),
                    min_value=0.0
                )
                boat_jetty = st.number_input(
                    "⛵ Boat Jetty", 
                    value=float(st.session_state.form_data.get("boat_jetty", 0)),
                    min_value=0.0
                )
                weekend_transport = st.number_input(
                    "🗓️ Uang Transport di tanggal Merah", 
                    value=float(st.session_state.form_data.get("weekend_transport", 0)),
                    min_value=0.0
                )
    
    # Calculate total
    total_cost = (hotel_cost + deposit + plane_cost + miscellaneous + airport_tax + 
                  ship_cost + train_cost + bus_cost + fuel_cost + toll_cost + 
                  taxi_cost + local_transport + boat_jetty + weekend_transport)
    
    # Summary
    with st.container():
        st.markdown("### 📊 Ringkasan")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("💰 Total Biaya", f"Rp {total_cost:,.0f}")
        with col2:
            duration = (end_date - start_date).days + 1 if end_date >= start_date else 1
            st.metric("📅 Durasi", f"{duration} hari")
        with col3:
            st.metric("🚢 Kapal", vessel_code or "Belum diisi")
    
    # Email and Generate
    st.markdown("### 📧 Pengiriman Laporan")
    email_rbd = st.text_input(
        "📬 Email Penerima RBD", 
        placeholder="contoh@email.com",
        help="Email yang akan menerima laporan RBD"
    )
    
    if st.button("🚀 Generate & Kirim RBD", type="primary", use_container_width=True):
        if not trip_id:
            st.error("❌ Silakan masukkan Trip ID untuk melanjutkan.")
            return
        
        if not email_rbd:
            st.error("❌ Silakan masukkan alamat email penerima.")
            return
        
        # Prepare data
        trip_data = {
            "trip_id": trip_id,
            "start_date": start_date,
            "end_date": end_date,
            "trip_purpose": trip_purpose,
            "vessel_code": vessel_code,
            "hotel_cost": hotel_cost,
            "deposit": deposit,
            "plane_cost": plane_cost,
            "miscellaneous": miscellaneous,
            "airport_tax": airport_tax,
            "ship_cost": ship_cost,
            "train_cost": train_cost,
            "bus_cost": bus_cost,
            "fuel_cost": fuel_cost,
            "toll_cost": toll_cost,
            "taxi_cost": taxi_cost,
            "local_transport": local_transport,
            "boat_jetty": boat_jetty,
            "weekend_transport": weekend_transport
        }
        
        generate_rbd_report(trip_data, email_rbd)

# ==============================================================================
# REPORT GENERATION FUNCTIONS
# ==============================================================================

def generate_inspection_report(vessel_name, imo, ship_type, callsign, place, survey_date, 
                               master, surveyor, activities, email_to_send):
    """Generate dan kirim laporan inspeksi"""
    
    if not email_to_send:
        st.error("❌ Silakan masukkan alamat email penerima.")
        return
    
    # Validasi input
    if not all([vessel_name, place, master, surveyor]):
        st.error("❌ Mohon lengkapi semua field yang diperlukan.")
        return
    
    try:
        with st.spinner('📥 Mengambil template dari GitHub...'):
            response = requests.get(TEMPLATE_INSPEKSI_URL)
            response.raise_for_status()
        template_file = io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Gagal mengambil template dari GitHub: {e}")
        return
    
    try:
        with st.spinner('📝 Membuat laporan inspeksi...'):
            doc = Document(template_file)
        
        # Replace placeholders
        replace_placeholder_everywhere(doc, "<<vessel_name>>", vessel_name)
        replace_placeholder_everywhere(doc, "<<type>>", ship_type)
        replace_placeholder_everywhere(doc, "<<place>>", f"{place}, {survey_date.strftime('%d %B %Y')}")
        replace_placeholder_everywhere(doc, "<<master>>", master)
        
        # Update surveyor jika berbeda dari default
        if surveyor != "Mohammad Fajar S":
            replace_placeholder_everywhere(doc, "Mohammad Fajar S", surveyor)
        
        # Add activities to table
        add_activities_to_table(doc, activities)
        
        # Save document
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        base_filename = f"{survey_date.strftime('%Y.%m.%d')} {vessel_name} Inspection Report"
        docx_filename = f"{base_filename}.docx"
        
        # PERBAIKAN: Pindahkan tombol download ke luar blok pengiriman email
        # agar tombol tetap muncul meskipun email gagal.
        st.success("✅ Laporan berhasil dibuat!")
        
        st.download_button(
            label="📄 Download Laporan (DOCX)",
            data=docx_buffer.getvalue(),
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
        # Send email
        with st.spinner('📧 Mengirim laporan melalui email...'):
            attachments_list = [(docx_buffer.getvalue(), docx_filename)]
            success = send_email_with_attachment(
                EMAIL_SENDER, EMAIL_PASSWORD, email_to_send, SMTP_SERVER, SMTP_PORT,
                f"Laporan Inspeksi: {vessel_name}",
                f"Terlampir laporan inspeksi kapal {vessel_name} tanggal {survey_date.strftime('%d %B %Y')}.\n\nDetail:\n- Kapal: {vessel_name}\n- Tipe: {ship_type}\n- IMO: {imo}\n- Call Sign: {callsign}\n- Master: {master}\n- Surveyor: {surveyor}\n- Lokasi: {place}\n- Jumlah Kegiatan: {len(activities)}",
                attachments_list
            )
        
        if success:
            st.success(f"✅ Laporan berhasil dikirim ke {email_to_send}!")
        
    except Exception as e:
        st.error(f"❌ Gagal membuat laporan inspeksi: {e}")

def generate_rbd_report(trip_data, email_rbd):
    """Generate dan kirim laporan RBD"""
    
    try:
        with st.spinner('📥 Mengambil template RBD dari GitHub...'):
            response = requests.get(TEMPLATE_RBD_URL)
            response.raise_for_status()
        template_file = io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Gagal mengambil template RBD dari GitHub: {e}")
        return
    
    try:
        with st.spinner('📊 Membuat laporan RBD...'):
            rbd_wb = load_workbook(template_file)
            ws = rbd_wb.active
            
            # --- PENAMBAHAN KODE ---
            # Mengisi tanggal mulai dan selesai ke sel D11 dan H11
            ws['D11'] = trip_data["start_date"]
            ws['H11'] = trip_data["end_date"]
            # -----------------------

            # Mengisi data ke spreadsheet
            ws['C13'] = trip_data["trip_purpose"]
            ws['F13'] = trip_data["vessel_code"]
            ws['K13'] = "DPA"
            
            # Mengisi biaya
            ws['N20'] = float(trip_data["hotel_cost"])
            ws['N22'] = float(trip_data["deposit"])
            ws['N24'] = float(trip_data["plane_cost"])
            ws['N26'] = float(trip_data["miscellaneous"])
            ws['N28'] = float(trip_data["airport_tax"])
            ws['N30'] = float(trip_data["ship_cost"])
            ws['N33'] = float(trip_data["train_cost"])
            ws['N36'] = float(trip_data["bus_cost"])
            ws['N39'] = float(trip_data["fuel_cost"])
            ws['N40'] = float(trip_data["toll_cost"])
            ws['N42'] = float(trip_data["taxi_cost"])
            ws['N46'] = float(trip_data["local_transport"])
            ws['N47'] = float(trip_data["boat_jetty"])
            ws['N52'] = float(trip_data["weekend_transport"])
            
            # Menghitung durasi
            duration = (trip_data["end_date"] - trip_data["start_date"]).days + 1
            ws['I15'] = duration
            ws['I16'] = trip_data["start_date"].strftime("%d %B %Y")
            ws['I17'] = trip_data["end_date"].strftime("%d %B %Y")
            
            # Mengisi tanggal
            ws['L60'] = trip_data["end_date"]
            
            output_buffer = io.BytesIO()
            rbd_wb.save(output_buffer)
            output_buffer.seek(0)
            
            rbd_filename = f"RBD {trip_data['trip_id']}.xlsx"
            
            # Calculate total for email
            total_cost = sum([
                trip_data["hotel_cost"], trip_data["deposit"], trip_data["plane_cost"],
                trip_data["miscellaneous"], trip_data["airport_tax"], trip_data["ship_cost"],
                trip_data["train_cost"], trip_data["bus_cost"], trip_data["fuel_cost"],
                trip_data["toll_cost"], trip_data["taxi_cost"], trip_data["local_transport"],
                trip_data["boat_jetty"], trip_data["weekend_transport"]
            ])
            
            # Send email
            with st.spinner('📧 Mengirim laporan RBD melalui email...'):
                attachments_list = [(output_buffer.getvalue(), rbd_filename)]
                success = send_email_with_attachment(
                    EMAIL_SENDER, EMAIL_PASSWORD, email_rbd, SMTP_SERVER, SMTP_PORT,
                    f"Rincian Biaya Dinas: {trip_data['trip_id']}",
                    f"Terlampir laporan rincian biaya perjalanan dinas dengan ID {trip_data['trip_id']}.\n\nDetail:\n- Tujuan: {trip_data['trip_purpose']}\n- Kapal: {trip_data['vessel_code']}\n- Periode: {trip_data['start_date'].strftime('%d %B %Y')} - {trip_data['end_date'].strftime('%d %B %Y')}\n- Durasi: {duration} hari\n- Total Biaya: Rp {total_cost:,.0f}",
                    attachments_list
                )
            
            if success:
                st.success(f"✅ Laporan RBD berhasil dikirim ke {email_rbd}!")
                
                # Save to MongoDB
                save_data = {
                    "trip_id": trip_data["trip_id"],
                    "start_date": trip_data["start_date"].strftime("%Y-%m-%d"),
                    "end_date": trip_data["end_date"].strftime("%Y-%m-%d"),
                    "trip_purpose": trip_data["trip_purpose"],
                    "vessel_code": trip_data["vessel_code"],
                    "hotel_cost": str(trip_data["hotel_cost"]),
                    "deposit": str(trip_data["deposit"]),
                    "plane_cost": str(trip_data["plane_cost"]),
                    "miscellaneous": str(trip_data["miscellaneous"]),
                    "airport_tax": str(trip_data["airport_tax"]),
                    "ship_cost": str(trip_data["ship_cost"]),
                    "train_cost": str(trip_data["train_cost"]),
                    "bus_cost": str(trip_data["bus_cost"]),
                    "fuel_cost": str(trip_data["fuel_cost"]),
                    "toll_cost": str(trip_data["toll_cost"]),
                    "taxi_cost": str(trip_data["taxi_cost"]),
                    "local_transport": str(trip_data["local_transport"]),
                    "boat_jetty": str(trip_data["boat_jetty"]),
                    "weekend_transport": str(trip_data["weekend_transport"])
                }
                save_to_mongodb(save_data)
                
                # Tampilkan download button
                st.download_button(
                    label="📄 Download Laporan RBD (XLSX)",
                    data=output_buffer.getvalue(),
                    file_name=rbd_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary"
                )
            
    except Exception as e:
        st.error(f"❌ Gagal membuat laporan RBD: {e}")

def save_rbd_data(trip_id):
    """Save RBD data dari form saat ini"""
    if not trip_id:
        st.error("❌ Trip ID tidak boleh kosong untuk menyimpan data.")
        return
    
    # Get current form values
    form_data = st.session_state.form_data
    trip_data = {
        "trip_id": trip_id,
        "start_date": form_data.get("start_date", ""),
        "end_date": form_data.get("end_date", ""),
        "trip_purpose": form_data.get("trip_purpose", ""),
        "vessel_code": form_data.get("vessel_code", ""),
        "hotel_cost": form_data.get("hotel_cost", "0"),
        "deposit": form_data.get("deposit", "0"),
        "plane_cost": form_data.get("plane_cost", "0"),
        "miscellaneous": form_data.get("miscellaneous", "0"),
        "airport_tax": form_data.get("airport_tax", "0"),
        "ship_cost": form_data.get("ship_cost", "0"),
        "train_cost": form_data.get("train_cost", "0"),
        "bus_cost": form_data.get("bus_cost", "0"),
        "fuel_cost": form_data.get("fuel_cost", "0"),
        "toll_cost": form_data.get("toll_cost", "0"),
        "taxi_cost": form_data.get("taxi_cost", "0"),
        "local_transport": form_data.get("local_transport", "0"),
        "boat_jetty": form_data.get("boat_jetty", "0"),
        "weekend_transport": form_data.get("weekend_transport", "0")
    }
    
    if save_to_mongodb(trip_data):
        st.session_state.form_data = trip_data

# ==============================================================================
# MAIN APPLICATION
# ==============================================================================

def main():
    """Main application function"""
    
    # Page configuration
    st.set_page_config(
        page_title="Aplikasi Laporan DPA GLS",
        page_icon="🚢",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS untuk tampilan yang lebih bagus
    st.markdown("""
    <style>
    .main {
        padding-top: 2rem;
    }
    
    .stAlert {
        border-radius: 10px;
    }
    
    .stButton > button {
        border-radius: 20px;
        border: none;
        padding: 0.5rem 1rem;
        font-weight: 600;
        transition: all 0.3s;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
    }
    
    .stSelectbox > div > div {
        border-radius: 10px;
    }
    
    .stTextInput > div > div {
        border-radius: 10px;
    }
    
    .stNumberInput > div > div {
        border-radius: 10px;
    }
    
    .stTextArea > div > div {
        border-radius: 10px;
    }
    
    .metric-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .header-container {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    .card {
        background: white;
        padding: 1.5rem;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin: 1rem 0;
        border: 1px solid #e1e5e9;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Initialize session state
    init_session_state()
    
    # Header
    st.markdown("""
    <div class="header-container">
        <h1>🚢 Aplikasi Laporan DPA GLS</h1>
        <p>Sistem Terintegrasi untuk Laporan Inspeksi Kapal dan Rincian Biaya Dinas</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Render sidebar
    render_sidebar()
    
    # Main content based on selected report type
    if st.session_state.report_type == "Laporan Inspeksi":
        render_inspection_form()
    else:
        render_rbd_form()
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>© 2025 DPA GLS - Aplikasi Laporan</p>
        <p>Dibuat dengan ❤️ menggunakan Hati</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
